"""
Generate IEEE-formatted DOCX from the paper LaTeX source.
Run: python3 generate_docx.py
Output: liver_cirrhosis_paper.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── helpers ────────────────────────────────────────────────────────────────────

def clean_latex(text):
    """Strip common LaTeX commands, leaving readable prose."""
    # citations → [n]
    text = re.sub(r'\\cite\{[^}]*\}', '', text)
    # bold
    text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
    # italic / textit
    text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
    # math mode — keep content, remove $ signs
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    # common accented chars
    text = text.replace(r"Fr\'{e}chet", 'Fréchet')
    text = text.replace(r"\'{e}", 'é')
    text = text.replace(r"\c{c}", 'ç')
    text = text.replace(r"\"{o}", 'ö')
    text = text.replace(r"\"{u}", 'ü')
    text = text.replace(r"\'{a}", 'á')
    text = text.replace(r"Gon\c{c}alves", 'Gonçalves')
    text = text.replace(r"Vega-M\'{a}rquez", 'Vega-Márquez')
    text = text.replace(r"K\"{u}\c{c}\"{u}kak\c{c}al{\i}", 'Küçükakçalı')
    text = text.replace(r"R\"{o}chner", 'Röchner')
    # section refs
    text = re.sub(r'Section~\\ref\{[^}]*\}', 'the relevant section', text)
    text = re.sub(r'Fig\.~\\ref\{[^}]*\}', '[Figure]', text)
    text = re.sub(r'Table~\\ref\{[^}]*\}', '[Table]', text)
    text = re.sub(r'Tables~\\ref\{[^}]*\}\s+and~\\ref\{[^}]*\}', '[Tables]', text)
    # misc LaTeX
    text = re.sub(r'\\ref\{[^}]*\}', '', text)
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    text = re.sub(r'\\IEEEmembership\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\mathbf\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\boldsymbol\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\mathcal\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\text\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\left', '', text)
    text = re.sub(r'\\right', '', text)
    text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z!,. ]+', ' ', text)
    text = re.sub(r'[{}]', '', text)
    text = re.sub(r'~', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def set_font(run, size_pt=10, bold=False, italic=False, color=None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Times New Roman'
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text, style='Normal', size=10, bold=False,
                  italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(text)
    set_font(run, size_pt=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 12, 2: 11, 3: 10}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper() if level == 1 else text)
    set_font(run, size_pt=sizes.get(level, 10), bold=True)
    return p


def add_figure_placeholder(doc, caption, fig_num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'[INSERT FIGURE {fig_num} HERE]')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.name = 'Times New Roman'

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(2)
    cap_p.paragraph_format.space_after = Pt(10)
    cap_run = cap_p.add_run(f'Fig. {fig_num}. {caption}')
    cap_run.font.size = Pt(9)
    cap_run.font.name = 'Times New Roman'


def add_table_from_data(doc, caption, headers, rows, table_num):
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(8)
    cap_p.paragraph_format.space_after = Pt(3)
    cap_run = cap_p.add_run(f'TABLE {table_num}\n{caption.upper()}')
    cap_run.font.size = Pt(9)
    cap_run.font.bold = True
    cap_run.font.name = 'Times New Roman'

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        trow = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = trow.cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing after table


# ── build document ──────────────────────────────────────────────────────────────

doc = Document()

# Page margins (IEEE-like: 1 inch all sides)
section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# ── TITLE ──────────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(8)
title_run = title_p.add_run(
    'Evaluating Predictive Utility of Synthetic Liver Cirrhosis Data in IoMT: '
    'A Comparative Study of GAN, CTGAN, and TVAE with Consensus-Based Quality Filtering'
)
set_font(title_run, size_pt=16, bold=True)

# ── AUTHORS ─────────────────────────────────────────────────────────────────────
auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.paragraph_format.space_after = Pt(2)
auth_run = auth_p.add_run('Michael Udousoro, Francis Anyaka, and Mfon Ekpo, Senior Member, IEEE')
set_font(auth_run, size_pt=11, italic=True)

aff_p = doc.add_paragraph()
aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff_p.paragraph_format.space_after = Pt(2)
aff_run = aff_p.add_run('[Department of Computer Science, Your University, City, Country]')
set_font(aff_run, size_pt=10)

note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
note_p.paragraph_format.space_after = Pt(14)
note_run = note_p.add_run(
    '*** Fill in your institutional affiliations and email addresses before submission ***'
)
set_font(note_run, size_pt=9, italic=True, color=(180, 0, 0))

# ── ABSTRACT ───────────────────────────────────────────────────────────────────
add_heading(doc, 'Abstract', level=1)
abstract_text = (
    "Expanding rare clinical datasets for Internet of Medical Things (IoMT) applications is "
    "constrained by patient privacy regulations and disease rarity. Most existing work evaluates "
    "synthetic data through distributional similarity metrics, without directly testing whether "
    "augmentation improves prediction on real patients. This paper investigates that question "
    "using the Mayo Clinic Primary Biliary Cirrhosis (PBC) dataset, which yields 276 complete "
    "records after removing incomplete cases from 418 patients. Three generative architectures "
    "are trained from scratch in TensorFlow 2.15: a Vanilla GAN, a Conditional Tabular GAN "
    "(CTGAN), and a Tabular Variational Autoencoder (TVAE). Generated records pass through an "
    "IQR-based physiological plausibility filter, followed by a multi-method consensus voting "
    "procedure that retains only records independently corroborated by all three models. "
    "Distributional fidelity is assessed through tabular Fréchet Inception Distance, "
    "Kolmogorov-Smirnov tests, and Jensen-Shannon divergence; privacy risk through "
    "nearest-neighbour distance analysis. Predictive utility is evaluated by training Random "
    "Forest, Gradient Boosting, and Logistic Regression classifiers under four augmentation "
    "scenarios and testing on 83 held-out real patients, with 5-fold cross-validation, "
    "1000-resample bootstrap confidence intervals, and McNemar's exact test. CTGAN achieves "
    "the best distributional fidelity at FID 0.075 with a near-duplicate rate of 9.0 percent. "
    "No augmentation scenario produces a statistically significant improvement in held-out "
    "accuracy, confirming that distributional similarity does not guarantee predictive gain on "
    "real patients. Consensus augmentation reduces Random Forest cross-validation AUC variance "
    "four-fold, from 0.054 to 0.014, demonstrating that cross-model agreement selects synthetic "
    "patients that stabilise classifiers across training subsets. This stability property has "
    "direct relevance for reliable clinical decision support in IoMT deployments where consistent "
    "performance matters across patient subpopulations."
)
add_paragraph(doc, abstract_text, size=10)

kw_p = doc.add_paragraph()
kw_p.paragraph_format.space_before = Pt(4)
kw_p.paragraph_format.space_after = Pt(12)
kw_bold = kw_p.add_run('Index Terms — ')
set_font(kw_bold, size_pt=10, bold=True, italic=True)
kw_text = kw_p.add_run(
    'Synthetic data generation, generative adversarial networks, variational autoencoder, '
    'liver cirrhosis, Internet of Medical Things, predictive utility, consensus voting, '
    'IQR filtering, patient privacy.'
)
set_font(kw_text, size_pt=10, italic=True)


# ── SECTION I: INTRODUCTION ────────────────────────────────────────────────────
add_heading(doc, 'I. Introduction', level=1)

add_paragraph(doc, (
    "The Internet of Medical Things has transformed how clinical biomarkers are collected, enabling "
    "continuous monitoring of liver enzymes, bilirubin concentrations, and albumin levels from wearable "
    "and implanted devices that feed directly into clinical decision support systems [1], [2]. As these "
    "systems grow more sophisticated, the demand for large, well-annotated patient datasets to train "
    "predictive models grows alongside them. The tension between that demand and the reality of clinical "
    "data availability is one of the defining challenges in medical machine learning. Privacy regulations, "
    "including GDPR and HIPAA, restrict what institutions can share, while rare conditions produce datasets "
    "that are inherently small. The Mayo Clinic Primary Biliary Cirrhosis trial is a concrete example: it "
    "enrolled 312 patients over more than a decade, and once records with incomplete measurements are "
    "removed, only 276 patients remain available for analysis [3]. Training robust classifiers on a dataset "
    "of that size is genuinely difficult."
))

add_paragraph(doc, (
    "Synthetic data generation through generative artificial intelligence has emerged as the most "
    "principled response to this situation. Rather than sharing real patient records, researchers can "
    "generate synthetic patients that statistically resemble the real population while not corresponding "
    "to any individual, enabling collaborative research without breaching privacy [4], [5]. Generative "
    "Adversarial Networks and Variational Autoencoders have demonstrated genuine promise across medical "
    "imaging, physiological time series, and tabular clinical records [6], [7], [8]. The difficulty is "
    "that the field has developed a habit of evaluating synthetic data through distributional lenses, "
    "measuring how similar synthetic samples are to real ones using feature-level statistics and "
    "aggregate distance metrics such as Fréchet Inception Distance or Jensen-Shannon divergence. These "
    "are sensible measures of generation quality, but they answer a different question from the one that "
    "matters most clinically. The relevant operational question is whether augmenting a training set with "
    "synthetic patients actually produces a model that predicts real patient outcomes more accurately. "
    "Whether distributional similarity translates into predictive utility has received far less systematic "
    "attention, particularly for small rare tabular clinical datasets in IoMT settings [9], [10]."
))

add_paragraph(doc, (
    "This paper addresses that gap through a controlled experimental pipeline applied to the PBC "
    "complete-case dataset. Three deep generative models are trained from scratch in TensorFlow 2.15: "
    "a Vanilla GAN, a Conditional Tabular GAN that conditions generation on the patient outcome label, "
    "and a Tabular Variational Autoencoder. Two pre-augmentation quality controls are then applied. The "
    "first is an IQR-based physiological plausibility filter that removes synthetic records containing "
    "biomarker values outside Tukey fence bounds computed from real training patients. The second is a "
    "multi-method consensus voting mechanism that identifies a high-confidence synthetic subset agreed "
    "upon independently by all three generative models, based on the premise that a synthetic patient "
    "produced in the same region of feature space by three structurally different architectures is more "
    "likely to reflect a genuine distributional pattern than one produced by any single model."
))

add_paragraph(doc, (
    "The contributions of this work are as follows. First, three deep generative architectures are "
    "compared for tabular liver cirrhosis data in an IoMT setting, all built from scratch, allowing a "
    "fair head-to-head assessment of adversarial and variational approaches on the same dataset. Second, "
    "an IQR-based plausibility filter and a multi-method consensus voting mechanism are introduced as "
    "pre-augmentation quality controls. Third, a predictive utility evaluation framework is applied that "
    "benchmarks each augmentation strategy against both a real-data baseline and a SMOTE oversampling "
    "baseline, using 5-fold cross-validation, bootstrap confidence intervals, and McNemar's exact test "
    "on a held-out real test set."
))

add_paragraph(doc, (
    "The paper proceeds as follows. Section II reviews relevant literature. Section III describes the "
    "dataset, generative architectures, filtering pipeline, and experimental design. Section IV presents "
    "results. Section V draws conclusions."
))


# ── SECTION II: RELATED WORK ───────────────────────────────────────────────────
add_heading(doc, 'II. Related Work', level=1)

add_heading(doc, 'A. Synthetic Data Generation for Healthcare', level=2)
add_paragraph(doc, (
    "Synthetic data generation has attracted sustained research interest as a means of making healthcare "
    "datasets more widely available without compromising patient privacy [4], [5], [11]. Gonçalves et al. "
    "conducted a systematic review of generation and evaluation methods, establishing that downstream "
    "utility, meaning the capacity of synthetic data to substitute for real data in analytical tasks, "
    "deserves equal attention alongside privacy preservation [4]. Gonzales et al. surveyed clinical "
    "applications across oncology, cardiology, and infectious disease, identifying tabular electronic "
    "health records as the domain with the greatest unmet need for robust generation methods [5]."
))
add_paragraph(doc, (
    "Classical augmentation approaches have well-documented limitations in this context. SMOTE addresses "
    "class imbalance by interpolating between existing minority-class training points and cannot generate "
    "samples outside the convex hull of observed data [12], [13]. Rule-based and parametric generators "
    "offer transparency but fail to capture the non-linear dependencies characterising clinical biomarker "
    "profiles. Deep generative models learn these dependencies directly from data and have become the "
    "dominant approach for tabular clinical synthesis [6], [7]."
))

add_heading(doc, 'B. GAN Architectures for Tabular Clinical Data', level=2)
add_paragraph(doc, (
    "Generative Adversarial Networks train a generator against a discriminator through adversarial "
    "feedback [14]. In clinical settings, Che et al. showed that GAN-augmented training data can improve "
    "risk prediction from electronic health records [8], and Choi et al. proposed medGAN for generating "
    "multi-label discrete patient records [15]. Standard GAN architectures are prone to mode collapse and "
    "training instability when applied to heterogeneous tabular data combining continuous biomarkers, "
    "binary indicators, and ordinal staging variables [16], [10]. The Conditional Tabular GAN addresses "
    "these limitations by conditioning both generator and discriminator on a class label, ensuring "
    "minority classes receive proportional representation [14], [17]. The Tabular Variational Autoencoder "
    "takes a complementary approach, optimising a probabilistic lower bound on the data likelihood. Its "
    "encoder maps patient records into a continuous latent space, while its decoder reconstructs records "
    "from samples drawn from that distribution, generating new patients at inference time by sampling "
    "from a standard normal prior [7]."
))

add_heading(doc, 'C. Synthetic Data for IoMT', level=2)
add_paragraph(doc, (
    "IoMT infrastructures produce continuous streams of physiological measurements that clinical decision "
    "support systems must process reliably, requiring large labelled training sets [1], [2]. Meidan et al. "
    "showed that GAN-generated data improved anomaly detection in IoMT systems [18]. D'Amico et al. "
    "demonstrated that AI-generated synthetic cohorts can accelerate research in haematological conditions "
    "where real patient numbers are similarly limited [19]. Kim et al. showed that synthetic augmentation "
    "improves survival prediction in early-onset colorectal cancer, a rare-disease tabular setting "
    "structurally similar to PBC [20]. Tucker et al. established that high-fidelity synthetic patient "
    "data can substitute for real data in evaluating healthcare software, provided quality is rigorously "
    "assessed [21]."
))

add_heading(doc, 'D. From Distributional Fidelity to Predictive Utility', level=2)
add_paragraph(doc, (
    "The dominant evaluation framework measures how closely generated samples resemble real data through "
    "marginal distributions and aggregate distance metrics [6], [13], [22]. Espinosa and Figueira found "
    "only weak correlation between distributional and utility metrics across a range of datasets [10]. "
    "Foraker et al. showed that analytical results from synthetic data do not reliably replicate real-data "
    "findings even when distributional similarity is high [23]. Kababji et al. concluded that the ability "
    "to reproduce hazard ratios from real patient data is a more clinically meaningful benchmark than FID "
    "alone [9]. Kaabachi et al. identified the absence of held-out real-patient evaluation as a critical "
    "and recurring gap in the medical synthetic data literature [24]. The present work is designed "
    "specifically to fill that gap."
))

add_heading(doc, 'E. Multi-Model Ensemble Strategies and the Positioning of This Work', level=2)
add_paragraph(doc, (
    "Jiang et al. proposed a multi-source data augmentation framework for industrial fault diagnosis in "
    "which several GAN models are trained independently and their outputs pooled through cooperative and "
    "competitive mechanisms [25]. Their central observation is that individual GAN models converge to "
    "different regions of the data manifold, so combining outputs produces a more representative "
    "augmented set than any single generator. The consensus voting mechanism introduced here shares this "
    "intuition but departs from it in two important respects: three structurally distinct architectures "
    "are used rather than multiple instances of the same one, and acceptance requires cross-architecture "
    "agreement rather than simply pooling all outputs."
))
add_paragraph(doc, (
    "Within the IoMT domain, Pandey et al. introduced MF-CGAN, a multifeature conditional GAN published "
    "in this journal that generates synthetic network traffic and health metrics by conditioning on "
    "multiple concurrent variables [26]. The present work complements MF-CGAN in the clinical direction: "
    "the data consists of continuous laboratory biomarkers, ordinal disease staging, and binary clinical "
    "indicators measured at a single visit. The evaluation criterion differs as well: classifiers are "
    "trained on augmented data and tested exclusively on real held-out patients. The consensus voting "
    "mechanism introduced here has no analogue in MF-CGAN, and the results reveal that its primary "
    "benefit is not improved peak accuracy but substantially reduced model variance across "
    "cross-validation folds."
))


# ── SECTION III: METHODOLOGY ───────────────────────────────────────────────────
add_heading(doc, 'III. Methodology', level=1)

add_heading(doc, 'A. Dataset and Preprocessing', level=2)
add_paragraph(doc, (
    "The dataset used in this study is the Mayo Clinic Primary Biliary Cirrhosis trial obtained from "
    "the UCI Machine Learning Repository [3]. The raw file contains 418 patient records and 20 columns. "
    "The patient identifier column carries no clinical information and is discarded. The remaining 19 "
    "columns include 11 continuous laboratory measurements (N_Days, Age, Bilirubin, Cholesterol, "
    "Albumin, Copper, Alk_Phos, SGOT, Tryglicerides, Platelets, and Prothrombin); five binary clinical "
    "indicators (Sex, Drug, Ascites, Hepatomegaly, and Spiders); two ordinal variables, Edema coded as "
    "0 for no oedema, 1 for slight or suspected oedema, and 2 for confirmed oedema, and Stage coded "
    "from 1 through 4; and the binary survival outcome Status."
))
add_paragraph(doc, (
    "Patients 313 through 418 belong to an observational cohort missing between 9 and 10 laboratory "
    "measurements each. Complete-case analysis is applied: any row containing at least one missing value "
    "is removed, yielding 276 complete patient records. Categorical columns are numerically encoded: Sex "
    "is Female = 0, Male = 1; Drug is Placebo = 0, D-penicillamine = 1; Ascites, Hepatomegaly, and "
    "Spiders are No = 0, Yes = 1; Status is censored or transplanted = 0, deceased = 1."
))
add_paragraph(doc, (
    "[Figure 1 shows the Pearson correlation matrix across all 276 complete cases. Bilirubin shows "
    "strong positive correlations with Copper (r = 0.46), SGOT (r = 0.42), Cholesterol (r = 0.39), "
    "and Prothrombin (r = 0.33), reflecting the interlinked nature of hepatic synthetic and excretory "
    "function in PBC. N_Days is negatively correlated with Bilirubin (r = -0.43) and Copper (r = "
    "-0.36), consistent with more severe biochemical profiles being associated with shorter survival.]"
))
add_figure_placeholder(doc,
    "Pearson correlation matrix computed across all 276 complete cases for the 11 continuous clinical "
    "features. Strong positive associations are visible between Bilirubin and Copper (r = 0.46), "
    "Bilirubin and SGOT (r = 0.42), and Bilirubin and Cholesterol (r = 0.39).",
    fig_num=1)
add_paragraph(doc, (
    "The 276 records are divided into a training set of 193 patients and a held-out test set of 83 "
    "patients using stratified random splitting with seed 42, preserving class balance in both "
    "partitions. The training set contains 115 survivors and 78 deceased patients. A MinMaxScaler is "
    "fitted on the training set alone and maps all 19 features to [0, 1]. The overall pipeline is "
    "illustrated in [Figure 2]."
))
add_figure_placeholder(doc,
    "End-to-end synthetic data pipeline. The raw 418-patient dataset passes through a complete-case "
    "filter and preprocessing stage before the 193-patient training set feeds three parallel generative "
    "models. Each model's 500 generated records pass through an IQR plausibility filter before all "
    "three filtered pools enter the consensus voting stage.",
    fig_num=2)

add_heading(doc, 'B. Generative Model Architectures', level=2)
add_paragraph(doc, (
    "All three models are implemented in TensorFlow 2.15 without third-party synthetic data libraries. "
    "Each model receives MinMaxScaled training data and produces samples in the same normalised range, "
    "which are inverse-transformed to recover original clinical units after generation."
))

p_gan = doc.add_paragraph()
p_gan.paragraph_format.space_before = Pt(4)
p_gan.paragraph_format.space_after = Pt(4)
p_gan.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
b1 = p_gan.add_run('Vanilla GAN. ')
set_font(b1, size_pt=10, bold=True)
r1 = p_gan.add_run(
    "The generator maps a 100-dimensional noise vector z drawn from a standard normal distribution "
    "through two dense hidden layers of 128 and 256 units with Leaky ReLU activations (alpha = 0.2) "
    "and batch normalisation (momentum = 0.8) to a sigmoid output of dimension 18. The discriminator "
    "maps each patient record through two dense layers of 256 and 128 units with Leaky ReLU activations "
    "and dropout (p = 0.3). Both networks use the Adam optimiser with learning rate 2e-4 and beta_1 = "
    "0.5 for 500 epochs at batch size 32. The adversarial objectives are: "
    "L_D = -E[log D(x)] - E[log(1 - D(G(z)))] and L_G = -E[log D(G(z))]."
)
set_font(r1, size_pt=10)

p_ctgan = doc.add_paragraph()
p_ctgan.paragraph_format.space_before = Pt(4)
p_ctgan.paragraph_format.space_after = Pt(4)
p_ctgan.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
b2 = p_ctgan.add_run('Conditional Tabular GAN (CTGAN). ')
set_font(b2, size_pt=10, bold=True)
r2 = p_ctgan.add_run(
    "The CTGAN extends the Vanilla GAN by conditioning both generator and discriminator on the patient "
    "outcome label y, supplied as a two-dimensional one-hot vector. The generator receives the "
    "concatenation [z || y] and the discriminator receives [x || y]. Training uses balanced "
    "mini-batches drawing equal numbers of survivors and deceased patients at each step. Architecture "
    "and optimiser settings match the Vanilla GAN, with training for 300 epochs."
)
set_font(r2, size_pt=10)

p_tvae = doc.add_paragraph()
p_tvae.paragraph_format.space_before = Pt(4)
p_tvae.paragraph_format.space_after = Pt(4)
p_tvae.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
b3 = p_tvae.add_run('Tabular Variational Autoencoder (TVAE). ')
set_font(b3, size_pt=10, bold=True)
r3 = p_tvae.add_run(
    "The encoder maps each patient record through two dense layers of 128 and 64 units with ReLU "
    "activations to a mean vector mu and log-variance vector log(sigma^2) of dimension 32. A latent "
    "sample is drawn using the reparameterisation trick: z = mu + epsilon * exp(0.5 * log(sigma^2)), "
    "where epsilon is drawn from a standard normal. The decoder maps z through two dense layers of 64 "
    "and 128 units with ReLU activations to a sigmoid output of dimension 18. The training objective is "
    "the Evidence Lower Bound (ELBO): L_TVAE = E[||x - x_hat||^2] + beta * KL(N(mu, sigma^2) || N(0, I)) "
    "with beta = 1.0, Adam optimiser at learning rate 1e-3, batch size 32, and 300 training epochs. "
    "Each model generates 500 synthetic records."
)
set_font(r3, size_pt=10)

add_paragraph(doc, (
    "[Figure 3] shows the training loss curves for all three models. The Vanilla GAN stabilises with "
    "generator loss near 0.8 and discriminator loss near 1.3 after approximately 100 epochs. The CTGAN "
    "experiences more volatile early training, with the generator loss peaking near 3.0 around epoch 30 "
    "before both networks converge by epoch 100. The TVAE ELBO loss drops sharply from 0.15 to "
    "approximately 0.08 by epoch 50 and then remains stable through the remaining 250 epochs."
))
add_figure_placeholder(doc,
    "Training loss curves for all three generative models. Left: Vanilla GAN generator and discriminator "
    "losses over 500 epochs. Centre: CTGAN losses over 300 epochs, showing a peak near epoch 30. "
    "Right: TVAE total ELBO loss over 300 epochs.",
    fig_num=3)

add_heading(doc, 'C. IQR-Based Physiological Plausibility Filter', level=2)
add_paragraph(doc, (
    "A Tukey fence filter is applied independently to each model's output. For each of the 11 continuous "
    "clinical features, the first and third quartiles, Q1 and Q3, are computed from the real training "
    "set alone. A synthetic record is retained only if every continuous feature value falls within "
    "[Q1 - 1.5 * IQR, Q3 + 1.5 * IQR]. A single out-of-range value disqualifies the entire record. "
    "Binary and ordinal features are exempt because they are snapped to valid integer ranges during "
    "post-processing."
))
add_paragraph(doc, (
    "[Figure 4] illustrates the filter effect on Bilirubin. For the GAN and CTGAN, the unfiltered "
    "distribution extends well beyond 15 mg/dL. The filter removes these records, leaving post-filter "
    "distributions that closely follow the real patient density near 0 to 5 mg/dL. The TVAE shows a "
    "qualitatively different pattern: the pre-filter and post-filter distributions are nearly identical "
    "because the smooth Gaussian latent space rarely extrapolates to extreme values, though the TVAE "
    "distribution is shifted rightward relative to real patients."
))
add_figure_placeholder(doc,
    "Effect of IQR-based Tukey fence filtering on the Bilirubin distribution for all three generative "
    "models. Grey bars show real patient density. Pink bars show unfiltered synthetic distribution; "
    "coloured bars show the post-filter distribution.",
    fig_num=4)
add_paragraph(doc, (
    "Applied to 500 generated records per model, the filter retained 226 GAN records (45.2%), "
    "210 CTGAN records (42.0%), and all 500 TVAE records (100%)."
))

add_heading(doc, 'D. Multi-Method Consensus Voting', level=2)
add_paragraph(doc, (
    "A consensus voting mechanism is applied that retains only records independently corroborated by "
    "all three generative models. This approach is related to the multi-source data augmentation "
    "strategy of Jiang et al. [25], who combine outputs from multiple independently trained GANs to "
    "reduce distribution skewness. The present work differs in that three structurally distinct "
    "architectures are used and acceptance requires cross-architecture agreement rather than simply "
    "pooling all outputs."
))
add_paragraph(doc, (
    "A StandardScaler is fitted on the union of all three filtered datasets. For each candidate record "
    "from one model, the minimum Euclidean distance to any record from each of the other two models is "
    "computed in standardised space. The candidate is accepted if both other models have a record within "
    "tolerance = 0.5 standardised units. This process repeats with each filtered dataset as the "
    "candidate pool in turn, and exact duplicates are removed from the final set."
))
add_paragraph(doc, (
    "[Figure 5] shows the breakdown of the 782 consensus records. GAN contributed 137 records (17.5%), "
    "CTGAN contributed 161 (20.6%), and TVAE contributed 484 (61.9%). The TVAE's dominant contribution "
    "reflects both its complete IQR retention and its broader coverage of feature space."
))
add_figure_placeholder(doc,
    "Source composition of the 782-record consensus set. Left: absolute record counts per generative "
    "method after consensus voting. Right: proportional contribution as a pie chart. TVAE contributes "
    "61.9%, GAN 17.5%, and CTGAN 20.6%.",
    fig_num=5)

add_heading(doc, 'E. Synthetic Data Quality Evaluation', level=2)
add_paragraph(doc, (
    "Distributional fidelity is assessed using a tabular proxy Fréchet Inception Distance computed on "
    "the 18-dimensional feature vectors: FID = ||mu_r - mu_s||^2 + Tr(Sigma_r + Sigma_s - "
    "2*(Sigma_r * Sigma_s)^(1/2)). Lower values indicate greater distributional alignment. "
    "Per-feature Kolmogorov-Smirnov tests, Jensen-Shannon divergence, Cohen's d, Shapiro-Wilk "
    "normality tests, and chi-square tests for categorical features provide a granular distributional "
    "profile. Privacy risk is assessed through nearest-neighbour distance analysis reporting the "
    "percentage of synthetic records within a standardised distance of 1.0 from any real training patient."
))

add_heading(doc, 'F. Predictive Utility Evaluation', level=2)
add_paragraph(doc, (
    "Predictive utility is evaluated by training Random Forest (RF), Gradient Boosting (GB), and "
    "Logistic Regression (LR) classifiers under four training scenarios and testing on the 83-patient "
    "held-out real test set:"
))
scenarios = [
    "Scenario A: 193 real patients only (baseline).",
    "Scenario B: 193 real + 210 IQR-filtered CTGAN records (403 total). CTGAN is selected here because it achieves the lowest FID among individual methods.",
    "Scenario C: 193 real + 782 consensus records (975 total).",
    "Scenario D: 193 real patients augmented with SMOTE oversampling (classical baseline).",
]
for s in scenarios:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(s)
    set_font(run, size_pt=10)

add_paragraph(doc, (
    "All classifiers use default scikit-learn hyperparameters. Evaluation metrics are Accuracy, F1, "
    "Precision, Recall, and AUC-ROC. 5-fold stratified cross-validation, 1000-resample bootstrap "
    "confidence intervals, and McNemar's exact test complement the single-split evaluation. All random "
    "seeds are fixed at 42."
))


# ── SECTION IV: RESULTS AND DISCUSSION ────────────────────────────────────────
add_heading(doc, 'IV. Results and Discussion', level=1)

add_heading(doc, 'A. Synthetic Data Retention and IQR Filtering', level=2)
add_paragraph(doc, (
    "Table I summarises IQR filter retention rates and FID scores. The GAN and CTGAN each reject more "
    "than half of their generated records; the TVAE retains all 500. The consensus procedure accepts "
    "782 records from the union of the three filtered pools."
))
add_table_from_data(doc,
    caption='Synthetic Data Retention and FID Scores',
    headers=['Method', 'Generated', 'Retained', 'Retention %', 'FID'],
    rows=[
        ['Vanilla GAN',  '500', '226', '45.2%',  '0.0972'],
        ['CTGAN',        '500', '210', '42.0%',  '0.0751'],
        ['TVAE',         '500', '500', '100.0%', '0.2194'],
        ['Consensus',    '—',   '782', '—',      '0.1116'],
    ],
    table_num='I')

add_heading(doc, 'B. Distributional Fidelity', level=2)
add_paragraph(doc, (
    "[Figure 6] shows the tabular FID scores for all four synthetic datasets. CTGAN achieves the lowest "
    "FID of 0.075, followed by the Vanilla GAN at 0.097. The TVAE produces an FID of 0.219 despite "
    "retaining all generated records, confirming that smooth variational generation reduces "
    "distributional sharpness. The consensus set scores 0.112."
))
add_figure_placeholder(doc,
    "Tabular Fréchet Inception Distance scores for all four synthetic datasets. CTGAN achieves the "
    "lowest FID at 0.075. The TVAE shows the highest FID at 0.219, reflecting over-smoothing. "
    "The consensus set scores 0.112.",
    fig_num=6)
add_paragraph(doc, (
    "[Figure 7] compares feature distributions for six representative clinical variables. Albumin and "
    "Prothrombin are well reproduced by GAN and CTGAN. Bilirubin is the most challenging feature: the "
    "TVAE distribution is noticeably shifted rightward relative to real patients, while GAN and CTGAN "
    "align better with the concentrated low-value peak."
))
add_figure_placeholder(doc,
    "Feature distribution comparison between real training data and all four synthetic methods for "
    "six representative clinical variables: Bilirubin, Albumin, Prothrombin, Cholesterol, Copper, "
    "and SGOT.",
    fig_num=7)
add_paragraph(doc, (
    "[Figure 8] compares the Pearson correlation structure of real patient data against the consensus "
    "synthetic dataset. The consensus set preserves the dominant pairwise associations, including the "
    "strong positive Bilirubin-Copper relationship and the negative N_Days-Bilirubin association, "
    "with moderate attenuation."
))
add_figure_placeholder(doc,
    "Pearson correlation structure comparison between real patient data (left) and consensus "
    "synthetic data (right). The consensus set preserves dominant clinical associations.",
    fig_num=8)
add_paragraph(doc, (
    "[Figure 9] shows the t-SNE projection of real patients and each synthetic method into two "
    "dimensions. The GAN and CTGAN filtered sets mix reasonably with the real patient manifold. The "
    "TVAE points spread more widely, with portions falling outside the real patient contours. The "
    "consensus set fills the central high-density region of the real patient manifold, confirming that "
    "cross-model agreement selects records from the most representative zone of feature space."
))
add_figure_placeholder(doc,
    "t-SNE two-dimensional projection of real patients and each synthetic method. GAN and CTGAN "
    "show reasonable overlap with the real patient manifold. The consensus set fills the central "
    "high-density region.",
    fig_num=9)
add_paragraph(doc, (
    "Per-feature KS tests reveal a consistent pattern. The filtered GAN shows no significant "
    "distributional difference from real training data on N_Days (KS = 0.095, p = 0.282), Age "
    "(KS = 0.067, p = 0.707), and Albumin (KS = 0.116, p = 0.111), but diverges significantly on "
    "Bilirubin (KS = 0.468), Alk_Phos (KS = 0.688), and SGOT (KS = 0.325). The TVAE diverges "
    "significantly on all 11 continuous features. The consensus set achieves substantially smaller "
    "KS statistics on the most challenging biomarkers: Bilirubin KS = 0.164 vs. GAN's 0.468; "
    "Alk_Phos KS = 0.181 vs. GAN's 0.688."
))

add_heading(doc, 'C. Privacy Risk Assessment', level=2)
add_paragraph(doc, (
    "[Figure 10] shows the privacy risk analysis. GAN and CTGAN distributions are centred well above "
    "the risk threshold, indicating that most adversarially generated records fall at a safe distance "
    "from any real patient. The TVAE and consensus distributions have substantial mass below the threshold."
))
add_figure_placeholder(doc,
    "Privacy disclosure risk analysis using nearest-neighbour distances in standardised feature "
    "space. Left: violin plots of distances from each synthetic record to its closest real training "
    "patient. Right: near-duplicate rates by method.",
    fig_num=10)

add_table_from_data(doc,
    caption='Privacy Risk Analysis',
    headers=['Method', 'Records', 'Mean Distance', 'Near-Duplicates', 'Rate'],
    rows=[
        ['GAN (filtered)',   '226', '2.391', '12',  '5.3%'],
        ['CTGAN (filtered)', '210', '2.286', '19',  '9.0%'],
        ['TVAE (filtered)',  '500', '1.591', '191', '38.2%'],
        ['Consensus',        '782', '1.754', '222', '28.4%'],
    ],
    table_num='II')

add_heading(doc, 'D. Predictive Utility on the Held-Out Test Set', level=2)
add_paragraph(doc, (
    "[Figure 11] shows the ROC curves for all three classifiers across Scenarios A, B, and C on the "
    "83-patient real held-out test set. The curves for all three scenarios lie close together for every "
    "classifier, with confidence band overlap across all scenarios. For Logistic Regression, Scenario C "
    "shows a modest upward shift relative to the baseline, with AUC improving from 0.829 to 0.855."
))
add_figure_placeholder(doc,
    "ROC curves by classifier and training scenario on the 83-patient real held-out test set. "
    "Random Forest (left), Gradient Boosting (centre), Logistic Regression (right). "
    "Gradient Boosting under Scenario A achieves the highest AUC at 0.878.",
    fig_num=11)
add_figure_placeholder(doc,
    "AUC scores across all three classifiers and training scenarios A, B, and C. All combinations "
    "exceed AUC = 0.8 comfortably.",
    fig_num=12)
add_figure_placeholder(doc,
    "F1-score heatmap across Scenarios A, B, and C for all three classifiers. Darker blue indicates "
    "higher F1. Scenario A achieves the highest F1 scores for Gradient Boosting and Random Forest "
    "at 0.785.",
    fig_num=13)

add_paragraph(doc, "Table III reports complete test-set metrics across all four scenarios.")
add_table_from_data(doc,
    caption='Test-Set Classification Performance (83 Real Held-Out Patients)',
    headers=['Scenario', 'Classifier', 'n_train', 'Accuracy', 'F1', 'Precision', 'Recall', 'AUC'],
    rows=[
        ['A: Baseline',       'Random Forest',       '193', '0.7831', '0.7848', '0.7901', '0.7831', '0.8479'],
        ['A: Baseline',       'Gradient Boosting',   '193', '0.7831', '0.7852', '0.7955', '0.7831', '0.8776'],
        ['A: Baseline',       'Logistic Regression', '193', '0.7711', '0.7735', '0.7940', '0.7711', '0.8291'],
        ['B: Real + CTGAN',   'Random Forest',       '403', '0.7590', '0.7616', '0.7782', '0.7590', '0.8418'],
        ['B: Real + CTGAN',   'Gradient Boosting',   '403', '0.7590', '0.7616', '0.7782', '0.7590', '0.8721'],
        ['B: Real + CTGAN',   'Logistic Regression', '403', '0.7590', '0.7616', '0.7782', '0.7590', '0.8455'],
        ['C: Real + Consensus', 'Random Forest',     '975', '0.7831', '0.7841', '0.7860', '0.7831', '0.8539'],
        ['C: Real + Consensus', 'Gradient Boosting', '975', '0.7349', '0.7377', '0.7617', '0.7349', '0.8321'],
        ['C: Real + Consensus', 'Logistic Regression','975', '0.7711', '0.7731', '0.7807', '0.7711', '0.8552'],
        ['D: Real + SMOTE',   'Random Forest',       '230', '0.7470', '0.7496', '0.7628', '0.7470', '0.8388'],
        ['D: Real + SMOTE',   'Gradient Boosting',   '230', '0.8072', '0.8087', '0.8140', '0.8072', '0.8661'],
        ['D: Real + SMOTE',   'Logistic Regression', '230', '0.7711', '0.7735', '0.7940', '0.7711', '0.8297'],
    ],
    table_num='III')

add_heading(doc, 'E. Cross-Validation Stability', level=2)
add_paragraph(doc, (
    "Table IV presents 5-fold stratified cross-validation AUC results. Consensus augmentation in "
    "Scenario C produces the highest mean CV AUC for both Random Forest (0.917) and Gradient Boosting "
    "(0.892). The standard deviation of CV AUC falls sharply: Random Forest variance drops from 0.054 "
    "in Scenario A to 0.014 in Scenario C, a four-fold reduction."
))
add_table_from_data(doc,
    caption='5-Fold Cross-Validation AUC (Mean ± Std)',
    headers=['Scenario', 'Random Forest', 'Gradient Boosting', 'Logistic Regression'],
    rows=[
        ['A: Baseline',         '0.884 ± 0.054', '0.854 ± 0.071', '0.852 ± 0.049'],
        ['B: Real + CTGAN',     '0.891 ± 0.028', '0.873 ± 0.032', '0.867 ± 0.034'],
        ['C: Real + Consensus', '0.917 ± 0.014', '0.892 ± 0.020', '0.873 ± 0.030'],
        ['D: Real + SMOTE',     '0.886 ± 0.045', '0.860 ± 0.063', '0.852 ± 0.052'],
    ],
    table_num='IV')

add_heading(doc, 'F. Bootstrap Confidence Intervals and McNemar\'s Test', level=2)
add_paragraph(doc, (
    "Tables V and VI report 95% bootstrap confidence intervals for test-set AUC and McNemar's exact "
    "test results respectively. All confidence intervals overlap substantially across scenarios, and "
    "all McNemar p-values exceed 0.28, confirming that no augmentation scenario produces a statistically "
    "significant improvement over the real-data baseline on the 83-patient test set."
))
add_table_from_data(doc,
    caption='95% Bootstrap Confidence Intervals for Test-Set AUC',
    headers=['Scenario', 'Classifier', 'AUC', '95% CI'],
    rows=[
        ['A: Baseline',         'RF',  '0.8479', '[0.754, 0.938]'],
        ['A: Baseline',         'GB',  '0.8776', '[0.788, 0.948]'],
        ['A: Baseline',         'LR',  '0.8291', '[0.728, 0.916]'],
        ['B: Real + CTGAN',     'RF',  '0.8418', '[0.742, 0.934]'],
        ['B: Real + CTGAN',     'GB',  '0.8721', '[0.776, 0.948]'],
        ['B: Real + CTGAN',     'LR',  '0.8455', '[0.746, 0.934]'],
        ['C: Real + Consensus', 'RF',  '0.8539', '[0.760, 0.933]'],
        ['C: Real + Consensus', 'GB',  '0.8321', '[0.731, 0.919]'],
        ['C: Real + Consensus', 'LR',  '0.8552', '[0.756, 0.935]'],
    ],
    table_num='V')

add_table_from_data(doc,
    caption="McNemar's Exact Test (Augmented vs. Baseline)",
    headers=['Classifier', 'Comparison', 'p-value', 'Significant'],
    rows=[
        ['Random Forest',       'A vs. B', '0.625', 'No'],
        ['Random Forest',       'A vs. C', '1.000', 'No'],
        ['Gradient Boosting',   'A vs. B', '0.688', 'No'],
        ['Gradient Boosting',   'A vs. C', '0.289', 'No'],
        ['Logistic Regression', 'A vs. B', '1.000', 'No'],
        ['Logistic Regression', 'A vs. C', '1.000', 'No'],
    ],
    table_num='VI')

add_heading(doc, 'G. Discussion', level=2)
add_paragraph(doc, (
    "The results offer a nuanced picture of what synthetic augmentation can and cannot achieve on a "
    "rare tabular clinical dataset in an IoMT context. On the held-out test set, no augmentation "
    "strategy produces a statistically significant improvement over training on real data alone. This "
    "is a meaningful finding rather than a failure: the 276-patient PBC dataset contains sufficient "
    "signal for Gradient Boosting to reach AUC = 0.878 without any augmentation, and adding synthetic "
    "patients that are distributional approximations of the real data does not substantially change "
    "what classifiers learn about the survival decision boundary when evaluated on 83 real patients."
))
add_paragraph(doc, (
    "The cross-validation results reveal a genuine benefit that single-split evaluation obscures. "
    "Consensus augmentation reduces Random Forest variance four-fold across folds, indicating that the "
    "additional synthetic patients stabilise the classifier across different subsets of the training data. "
    "In an IoMT deployment setting where a clinical decision support model must perform reliably across "
    "patients of different disease stages, ages, and treatment backgrounds, a model with lower "
    "cross-validation variance is more trustworthy than one that achieves peak performance on one "
    "partition by chance."
))
add_paragraph(doc, (
    "[Figure 14] shows feature importance rankings from Random Forest and Gradient Boosting trained "
    "under Scenario A. Both classifiers rank N_Days as the most important predictor (RF: 0.133, "
    "GB: 0.240), followed by Bilirubin (RF: 0.128, GB: 0.145). Prothrombin and Alk_Phos follow in "
    "both classifiers. Binary clinical indicators such as Ascites, Drug, Spiders, and Edema show low "
    "importance. This ranking is clinically coherent and validates the integrity of the preprocessing "
    "and modelling pipeline."
))
add_figure_placeholder(doc,
    "Feature importance rankings from Random Forest (left) and Gradient Boosting (right) trained "
    "on real data only under Scenario A. Both classifiers rank N_Days as the most important predictor, "
    "followed by Bilirubin.",
    fig_num=14)


# ── SECTION V: CONCLUSION ──────────────────────────────────────────────────────
add_heading(doc, 'V. Conclusion', level=1)

add_paragraph(doc, (
    "This study investigated whether deep generative synthetic data augmentation can improve liver "
    "cirrhosis survival prediction in an IoMT setting, using the Mayo Clinic PBC dataset as a "
    "controlled rare-disease tabular data scenario. Three generative architectures were trained from "
    "scratch and evaluated through a pipeline that goes beyond distributional fidelity to directly "
    "measure predictive utility on real held-out patients."
))
add_paragraph(doc, (
    "Two methodological contributions were introduced. An IQR-based physiological plausibility filter "
    "removes synthetic records with clinically implausible biomarker values before any downstream use, "
    "discarding between 54 and 58 percent of adversarially generated records. A multi-method consensus "
    "voting mechanism accepts only synthetic patients corroborated independently by all three generative "
    "architectures in standardised feature space."
))
add_paragraph(doc, (
    "Three empirical findings stand out. Among individual generative models, CTGAN produces the best "
    "distributional fidelity at FID = 0.075 and the most acceptable privacy profile at 9.0 percent "
    "near-duplicates, making it the most suitable model for privacy-constrained IoMT augmentation. "
    "On the 83-patient held-out test set, no augmentation scenario achieves a statistically significant "
    "improvement over the real-data baseline, confirming that distributional similarity does not "
    "guarantee predictive gain on real patients. Consensus augmentation reduces Random Forest "
    "cross-validation AUC variance four-fold, from 0.054 to 0.014, demonstrating that synthetic "
    "patients drawn from a cross-architecture agreement region stabilise classifiers in a way that "
    "random oversampling methods do not."
))
add_paragraph(doc, (
    "Future work should examine whether this stability effect translates into improved generalisation "
    "on external validation cohorts from different clinical sites. Incorporating differentially private "
    "training objectives into the TVAE architecture would address the memorisation risk identified in "
    "the privacy analysis. Extending the framework to longitudinal IoMT data streams, where temporal "
    "dependencies introduce additional generation challenges beyond those present in single-visit "
    "tabular records, represents a further natural direction for this line of research."
))


# ── REFERENCES ─────────────────────────────────────────────────────────────────
add_heading(doc, 'References', level=1)

references = [
    "[1] A. Guo, N. Mazumder, D. Ladner, and R. Foraker, \"Predicting mortality among patients with liver cirrhosis in electronic health records with machine learning,\" PLoS One, vol. 16, no. 8, p. e0256428, Aug. 2021.",
    "[2] J. Obeid, A. Khalifa, B. Xavier, H. Bou-Daher, and D. Rockey, \"An AI approach for identifying patients with cirrhosis,\" J. Clin. Gastroenterol., vol. 57, no. 1, pp. 82-88, Jan. 2023.",
    "[3] F. Kanwal, T. Taylor, J. Kramer, Y. Cao, D. Smith, A. Gifford et al., \"Development, validation, and evaluation of a simple machine learning model to predict cirrhosis mortality,\" JAMA Netw. Open, vol. 3, no. 11, p. e2023780, Nov. 2020.",
    "[4] R. S. Gonçalves et al., \"Generation and evaluation of synthetic patient data,\" BMC Med. Res. Methodol., vol. 20, no. 1, May 2020.",
    "[5] A. Gonzales, G. Guruswamy, and S. R. Smith, \"Synthetic data in health care: A narrative review,\" PLoS Digit. Health, vol. 2, no. 1, p. e0000082, Jan. 2023.",
    "[6] C. Yan et al., \"A multifaceted benchmarking of synthetic electronic health record generation models,\" Nat. Commun., vol. 13, no. 1, Dec. 2022.",
    "[7] G. Nikolentzos, M. Vazirgiannis, C. Xypolopoulos, M. Lingman, and E. G. Brandt, \"Synthetic electronic health records generated with variational graph autoencoders,\" NPJ Digit. Med., vol. 6, no. 1, Apr. 2023.",
    "[8] Z. Che, Y. Cheng, S. Zhai, Z. Sun, and Y. Liu, \"Boosting deep learning risk prediction with generative adversarial networks for electronic health records,\" in Proc. IEEE Int. Conf. Data Min., 2017, pp. 787-792.",
    "[9] S. E. Kababji et al., \"Evaluating the utility and privacy of synthetic breast cancer clinical trial data sets,\" JCO Clin. Cancer Inform., no. 7, p. e2300116, Nov. 2023.",
    "[10] J. Espinosa and S. Figueira, \"On the quality of synthetic generated tabular data,\" Mathematics, vol. 11, no. 15, p. 3278, Jul. 2023.",
    "[11] J. Guillaudeux et al., \"Patient-centric synthetic data generation, no reason to risk re-identification in biomedical data analysis,\" NPJ Digit. Med., vol. 6, no. 1, Mar. 2023.",
    "[12] J. Kumari and R. Siddharthan, \"MMM and MMMSynth: Clustering of heterogeneous tabular data, and synthetic data generation,\" PLoS One, vol. 19, no. 5, p. e0302271, May 2024.",
    "[13] S. Figueira and B. Vaz, \"Survey on synthetic data generation, evaluation methods and GANs,\" Mathematics, vol. 10, no. 15, p. 2733, Aug. 2022.",
    "[14] B. Vega-Márquez et al., \"Generation of synthetic data with conditional generative adversarial networks,\" Log. J. IGPL, vol. 28, no. 6, pp. 1098-1114, Dec. 2020.",
    "[15] E. Choi, S. Biswal, B. Malin, J. Duke, W. F. Stewart, and J. Sun, \"Generating multi-label discrete patient records using generative adversarial networks,\" arXiv, Mar. 2017.",
    "[16] S. Lenz et al., \"Deep generative models in DataSHIELD,\" BMC Med. Res. Methodol., vol. 21, no. 1, Apr. 2021.",
    "[17] D. Sharma, W. Lou, and W. Xu, \"PhylaGAN: Data augmentation through conditional GANs and autoencoders for improving disease prediction accuracy using microbiome data,\" Bioinformatics, vol. 40, no. 4, p. btae161, Apr. 2024.",
    "[18] Y. Meidan et al., \"N-BaIoT: Network-based detection of IoT botnet attacks using deep autoencoders,\" IEEE Pervasive Comput., vol. 17, no. 3, pp. 12-22, Jul.-Sep. 2018.",
    "[19] R. D'Amico et al., \"Synthetic data generation by artificial intelligence to accelerate research and precision medicine in hematology,\" JCO Clin. Cancer Inform., no. 7, p. e2300021, Jul. 2023.",
    "[20] J. Kim et al., \"Synthetic data improve survival status prediction models in early-onset colorectal cancer,\" JCO Clin. Cancer Inform., no. 8, p. e2300201, May 2024.",
    "[21] A. Tucker, Z. Wang, Y. Rotalinti, and P. Myles, \"Generating high-fidelity synthetic patient data for assessing machine learning healthcare software,\" NPJ Digit. Med., vol. 3, no. 1, Nov. 2020.",
    "[22] F. K. Dankar et al., \"A multi-dimensional evaluation of synthetic data generators,\" IEEE Access, vol. 10, pp. 12438-12454, 2022.",
    "[23] R. E. Foraker et al., \"Spot the difference: Comparing results of analyses from real patient data and synthetic derivatives,\" JAMIA Open, vol. 3, no. 4, pp. 557-566, Dec. 2020.",
    "[24] H. Kaabachi et al., \"A scoping review of privacy and utility metrics in medical synthetic data,\" medRxiv, Nov. 2023.",
    "[25] X. Jiang, J. Zheng, X. Zhuang, and Z. Ge, \"Ensemble data augmentation for imbalanced fault diagnosis,\" IEEE Trans. Instrum. Meas., vol. 72, p. 3528312, 2023.",
    "[26] C. Pandey, V. Tiwari, S. A. J. Francis, V. Pal, and D. S. Roy, \"MF-CGAN: Multifeature conditional GAN for synthetic data generation in Internet of Medical Things,\" IEEE Internet Things J., vol. 12, no. 10, pp. 13469-13485, May 2025.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run = p.add_run(ref)
    set_font(run, size_pt=9)


# ── SAVE ───────────────────────────────────────────────────────────────────────
output_path = '/Users/michaeludousoro/Desktop/liver_cirrhosis_article/paper/liver_cirrhosis_paper.docx'
doc.save(output_path)
print(f"DOCX saved: {output_path}")
